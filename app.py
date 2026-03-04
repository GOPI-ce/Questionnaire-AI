"""
QuestionnaireAI - Automated Questionnaire Answering Tool
Fictional Company: NovaMed Health (SaaS healthcare compliance platform)
"""

import os
import json
import sqlite3
import hashlib
import secrets
import re
from datetime import datetime
from functools import wraps
from flask import (
    Flask, render_template, request, redirect, url_for,
    session, jsonify, send_file, flash, g
)
from werkzeug.utils import secure_filename
import pypdf
import docx
from docx import Document
from docx.shared import Pt, RGBColor
from io import BytesIO

# ── Config ─────────────────────────────────────────────────────────────────
app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", secrets.token_hex(32))
app.config["UPLOAD_FOLDER"] = os.path.join(os.path.dirname(__file__), "uploads")
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024  # 16 MB
app.config["DB_PATH"] = os.path.join(os.path.dirname(__file__), "instance", "app.db")

ALLOWED_EXT = {"pdf", "docx", "txt"}
os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)
os.makedirs(os.path.dirname(app.config["DB_PATH"]), exist_ok=True)

# Load .env file
def load_dotenv():
    env_path = os.path.join(os.path.dirname(__file__), ".env")
    if os.path.exists(env_path):
        with open(env_path) as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith("#") and "=" in line:
                    key, _, val = line.partition("=")
                    os.environ.setdefault(key.strip(), val.strip())
load_dotenv()

OPENROUTER_API_KEY = os.environ.get("OPENROUTER_API_KEY", "")

@app.template_filter('fromjson')
def fromjson_filter(s):
    try:
        return json.loads(s)
    except Exception:
        return []

# ── DB helpers ──────────────────────────────────────────────────────────────
def get_db():
    if "db" not in g:
        g.db = sqlite3.connect(app.config["DB_PATH"])
        g.db.row_factory = sqlite3.Row
    return g.db

@app.teardown_appcontext
def close_db(e=None):
    db = g.pop("db", None)
    if db:
        db.close()

def init_db():
    db = sqlite3.connect(app.config["DB_PATH"])
    db.executescript("""
    CREATE TABLE IF NOT EXISTS users (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        email TEXT UNIQUE NOT NULL,
        password_hash TEXT NOT NULL,
        name TEXT NOT NULL,
        created_at TEXT DEFAULT (datetime('now'))
    );

    CREATE TABLE IF NOT EXISTS projects (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        name TEXT NOT NULL,
        created_at TEXT DEFAULT (datetime('now')),
        FOREIGN KEY(user_id) REFERENCES users(id)
    );

    CREATE TABLE IF NOT EXISTS reference_docs (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        project_id INTEGER NOT NULL,
        filename TEXT NOT NULL,
        content TEXT NOT NULL,
        created_at TEXT DEFAULT (datetime('now')),
        FOREIGN KEY(project_id) REFERENCES projects(id)
    );

    CREATE TABLE IF NOT EXISTS questionnaire_runs (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        project_id INTEGER NOT NULL,
        filename TEXT NOT NULL,
        status TEXT DEFAULT 'pending',
        created_at TEXT DEFAULT (datetime('now')),
        FOREIGN KEY(project_id) REFERENCES projects(id)
    );

    CREATE TABLE IF NOT EXISTS questions (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        run_id INTEGER NOT NULL,
        question_number INTEGER NOT NULL,
        question_text TEXT NOT NULL,
        answer TEXT,
        citations TEXT,
        confidence REAL DEFAULT 0,
        evidence_snippet TEXT,
        edited_answer TEXT,
        FOREIGN KEY(run_id) REFERENCES questionnaire_runs(id)
    );
    """)
    db.commit()
    db.close()

init_db()

# ── Auth helpers ────────────────────────────────────────────────────────────
def hash_password(pw):
    return hashlib.sha256(pw.encode()).hexdigest()

def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if "user_id" not in session:
            return redirect(url_for("login"))
        return f(*args, **kwargs)
    return decorated

def current_user():
    if "user_id" not in session:
        return None
    db = get_db()
    return db.execute("SELECT * FROM users WHERE id=?", (session["user_id"],)).fetchone()

# ── File parsing ────────────────────────────────────────────────────────────
def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXT

def extract_text(filepath):
    ext = filepath.rsplit(".", 1)[1].lower()
    if ext == "pdf":
        reader = pypdf.PdfReader(filepath)
        return "\n".join(p.extract_text() or "" for p in reader.pages)
    elif ext == "docx":
        doc = Document(filepath)
        return "\n".join(p.text for p in doc.paragraphs)
    elif ext == "txt":
        with open(filepath, encoding="utf-8", errors="ignore") as f:
            return f.read()
    return ""

def parse_questions(text):
    """Extract numbered questions from questionnaire text."""
    lines = text.splitlines()
    questions = []
    current = []
    number = 0

    patterns = [
        r"^\s*(\d+)[.)]\s+(.+)",
        r"^\s*Q(\d+)[.:)]\s*(.+)",
        r"^\s*Question\s+(\d+)[.:)]\s*(.+)",
    ]

    for line in lines:
        matched = False
        for pat in patterns:
            m = re.match(pat, line, re.IGNORECASE)
            if m:
                if current:
                    questions.append((number, " ".join(current).strip()))
                number = int(m.group(1))
                current = [m.group(2).strip()]
                matched = True
                break
        if not matched and current and line.strip():
            current.append(line.strip())

    if current:
        questions.append((number, " ".join(current).strip()))

    return questions

# ── RAG retrieval ────────────────────────────────────────────────────────────
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np

def chunk_text(text, chunk_size=300, overlap=50):
    words = text.split()
    chunks = []
    for i in range(0, len(words), chunk_size - overlap):
        chunk = " ".join(words[i:i + chunk_size])
        if chunk:
            chunks.append(chunk)
    return chunks

def retrieve_relevant_chunks(question, ref_docs, top_k=4):
    """TF-IDF retrieval across all reference doc chunks."""
    all_chunks = []
    chunk_meta = []  # (doc_id, filename, chunk_text)

    for doc in ref_docs:
        chunks = chunk_text(doc["content"])
        for c in chunks:
            all_chunks.append(c)
            chunk_meta.append({"doc_id": doc["id"], "filename": doc["filename"], "text": c})

    if not all_chunks:
        return []

    corpus = [question] + all_chunks
    try:
        vec = TfidfVectorizer(stop_words="english", max_features=5000)
        tfidf = vec.fit_transform(corpus)
        sims = cosine_similarity(tfidf[0:1], tfidf[1:]).flatten()
        top_indices = np.argsort(sims)[::-1][:top_k]
        results = []
        for idx in top_indices:
            if sims[idx] > 0.01:
                meta = chunk_meta[idx]
                meta["score"] = float(sims[idx])
                results.append(meta)
        return results
    except Exception:
        return []

# ── Claude API call ─────────────────────────────────────────────────────────
import urllib.request
import urllib.error

def call_claude(prompt, system=""):
    if not OPENROUTER_API_KEY:
        return "ERROR: OPENROUTER_API_KEY not set in .env file."

    # OpenRouter uses OpenAI-compatible format
    messages = []
    if system:
        messages.append({"role": "system", "content": system})
    messages.append({"role": "user", "content": prompt})

    payload = json.dumps({
        "model": "google/gemini-2.5-flash",
        "max_tokens": 1000,
        "messages": messages,
    }).encode()

    req = urllib.request.Request(
        "https://openrouter.ai/api/v1/chat/completions",
        data=payload,
        headers={
            "Content-Type": "application/json",
            "Authorization": f"Bearer {OPENROUTER_API_KEY}",
            "HTTP-Referer": "http://localhost:5000",
            "X-Title": "QuestionnaireAI",
        },
        method="POST"
    )
    try:
        with urllib.request.urlopen(req, timeout=60) as resp:
            data = json.loads(resp.read())
            return data["choices"][0]["message"]["content"]
    except urllib.error.HTTPError as e:
        body = e.read().decode()
        return f"API Error {e.code}: {body}"
    except Exception as ex:
        return f"Request failed: {ex}"

def generate_answer(question, chunks):
    if not chunks:
        return {
            "answer": "Not found in references.",
            "citations": [],
            "confidence": 0.0,
            "evidence_snippet": ""
        }

    context_parts = []
    for i, c in enumerate(chunks):
        context_parts.append(f"[Source {i+1}: {c['filename']}]\n{c['text']}")
    context = "\n\n---\n\n".join(context_parts)

    system = """You are a compliance analyst answering questionnaire questions using only the provided reference documents.
Rules:
- Answer based ONLY on the provided sources. 
- Be concise and factual.
- If the answer is not in the sources, say exactly: "Not found in references."
- Always cite which source(s) you used at the end as: [Source N: filename]
- Respond in JSON with keys: answer, citations (list of source filenames), confidence (0.0-1.0), evidence_snippet (short quote from source)
- confidence should reflect how directly the source material answers the question (1.0 = directly stated, 0.5 = inferred)"""

    prompt = f"""Reference Documents:
{context}

Question: {question}

Respond only with valid JSON."""

    raw = call_claude(prompt, system=system)

    # Parse JSON from response
    try:
        clean = re.sub(r"```json|```", "", raw).strip()
        data = json.loads(clean)
        return {
            "answer": data.get("answer", "Not found in references."),
            "citations": data.get("citations", []),
            "confidence": float(data.get("confidence", 0.5)),
            "evidence_snippet": data.get("evidence_snippet", "")
        }
    except Exception:
        # Fallback: treat raw as answer
        return {
            "answer": raw if raw else "Not found in references.",
            "citations": [c["filename"] for c in chunks[:2]],
            "confidence": 0.4,
            "evidence_snippet": chunks[0]["text"][:200] if chunks else ""
        }

# ── Routes ──────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    if "user_id" in session:
        return redirect(url_for("dashboard"))
    return redirect(url_for("login"))

@app.route("/signup", methods=["GET", "POST"])
def signup():
    if request.method == "POST":
        name = request.form.get("name", "").strip()
        email = request.form.get("email", "").strip().lower()
        pw = request.form.get("password", "")
        if not all([name, email, pw]):
            flash("All fields required.", "error")
            return render_template("signup.html")
        db = get_db()
        if db.execute("SELECT id FROM users WHERE email=?", (email,)).fetchone():
            flash("Email already registered.", "error")
            return render_template("signup.html")
        db.execute("INSERT INTO users(email,password_hash,name) VALUES(?,?,?)",
                   (email, hash_password(pw), name))
        db.commit()
        flash("Account created! Please log in.", "success")
        return redirect(url_for("login"))
    return render_template("signup.html")

@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        email = request.form.get("email", "").strip().lower()
        pw = request.form.get("password", "")
        db = get_db()
        user = db.execute("SELECT * FROM users WHERE email=? AND password_hash=?",
                          (email, hash_password(pw))).fetchone()
        if not user:
            flash("Invalid credentials.", "error")
            return render_template("login.html")
        session["user_id"] = user["id"]
        session["user_name"] = user["name"]
        return redirect(url_for("dashboard"))
    return render_template("login.html")

@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))

@app.route("/dashboard")
@login_required
def dashboard():
    db = get_db()
    user = current_user()
    projects = db.execute(
        "SELECT * FROM projects WHERE user_id=? ORDER BY created_at DESC",
        (user["id"],)
    ).fetchall()
    return render_template("dashboard.html", user=user, projects=projects)

@app.route("/project/new", methods=["POST"])
@login_required
def new_project():
    name = request.form.get("name", "").strip()
    if not name:
        flash("Project name required.", "error")
        return redirect(url_for("dashboard"))
    db = get_db()
    user = current_user()
    db.execute("INSERT INTO projects(user_id,name) VALUES(?,?)", (user["id"], name))
    db.commit()
    proj = db.execute("SELECT * FROM projects WHERE user_id=? ORDER BY id DESC LIMIT 1",
                      (user["id"],)).fetchone()
    return redirect(url_for("project", project_id=proj["id"]))

@app.route("/project/<int:project_id>")
@login_required
def project(project_id):
    db = get_db()
    user = current_user()
    proj = db.execute("SELECT * FROM projects WHERE id=? AND user_id=?",
                      (project_id, user["id"])).fetchone()
    if not proj:
        flash("Project not found.", "error")
        return redirect(url_for("dashboard"))
    ref_docs = db.execute("SELECT * FROM reference_docs WHERE project_id=? ORDER BY id",
                          (project_id,)).fetchall()
    runs = db.execute(
        "SELECT * FROM questionnaire_runs WHERE project_id=? ORDER BY created_at DESC",
        (project_id,)
    ).fetchall()
    return render_template("project.html", user=user, project=proj,
                           ref_docs=ref_docs, runs=runs)

@app.route("/project/<int:project_id>/upload_ref", methods=["POST"])
@login_required
def upload_ref(project_id):
    db = get_db()
    user = current_user()
    proj = db.execute("SELECT * FROM projects WHERE id=? AND user_id=?",
                      (project_id, user["id"])).fetchone()
    if not proj:
        return jsonify({"error": "Not found"}), 404

    files = request.files.getlist("files")
    added = []
    for f in files:
        if f and allowed_file(f.filename):
            fname = secure_filename(f.filename)
            path = os.path.join(app.config["UPLOAD_FOLDER"], fname)
            f.save(path)
            content = extract_text(path)
            db.execute(
                "INSERT INTO reference_docs(project_id,filename,content) VALUES(?,?,?)",
                (project_id, fname, content)
            )
            added.append(fname)
    db.commit()
    return jsonify({"added": added})

@app.route("/project/<int:project_id>/upload_questionnaire", methods=["POST"])
@login_required
def upload_questionnaire(project_id):
    db = get_db()
    user = current_user()
    proj = db.execute("SELECT * FROM projects WHERE id=? AND user_id=?",
                      (project_id, user["id"])).fetchone()
    if not proj:
        return jsonify({"error": "Not found"}), 404

    f = request.files.get("file")
    if not f or not allowed_file(f.filename):
        return jsonify({"error": "Invalid file"}), 400

    fname = secure_filename(f.filename)
    path = os.path.join(app.config["UPLOAD_FOLDER"], fname)
    f.save(path)
    content = extract_text(path)
    questions = parse_questions(content)

    if not questions:
        return jsonify({"error": "No questions found. Ensure questions are numbered (1. 2. etc)"}), 400

    db.execute(
        "INSERT INTO questionnaire_runs(project_id,filename,status) VALUES(?,?,?)",
        (project_id, fname, "pending")
    )
    db.commit()
    run = db.execute("SELECT * FROM questionnaire_runs WHERE project_id=? ORDER BY id DESC LIMIT 1",
                     (project_id,)).fetchone()
    run_id = run["id"]

    for num, qtext in questions:
        db.execute(
            "INSERT INTO questions(run_id,question_number,question_text) VALUES(?,?,?)",
            (run_id, num, qtext)
        )
    db.commit()
    return jsonify({"run_id": run_id, "question_count": len(questions)})

@app.route("/run/<int:run_id>/generate", methods=["POST"])
@login_required
def generate(run_id):
    db = get_db()
    user = current_user()
    run = db.execute("SELECT * FROM questionnaire_runs WHERE id=?", (run_id,)).fetchone()
    if not run:
        return jsonify({"error": "Run not found"}), 404

    proj = db.execute("SELECT * FROM projects WHERE id=? AND user_id=?",
                      (run["project_id"], user["id"])).fetchone()
    if not proj:
        return jsonify({"error": "Unauthorized"}), 403

    # Optionally regenerate only selected question ids
    q_ids = request.json.get("question_ids") if request.is_json else None

    ref_docs_raw = db.execute(
        "SELECT * FROM reference_docs WHERE project_id=?", (run["project_id"],)
    ).fetchall()
    ref_docs = [{"id": r["id"], "filename": r["filename"], "content": r["content"]}
                for r in ref_docs_raw]

    if q_ids:
        qs = db.execute(
            f"SELECT * FROM questions WHERE run_id=? AND id IN ({','.join('?'*len(q_ids))})",
            [run_id] + q_ids
        ).fetchall()
    else:
        qs = db.execute("SELECT * FROM questions WHERE run_id=?", (run_id,)).fetchall()

    db.execute("UPDATE questionnaire_runs SET status='generating' WHERE id=?", (run_id,))
    db.commit()

    results = []
    for q in qs:
        chunks = retrieve_relevant_chunks(q["question_text"], ref_docs)
        result = generate_answer(q["question_text"], chunks)
        db.execute("""
            UPDATE questions SET
                answer=?, citations=?, confidence=?, evidence_snippet=?, edited_answer=NULL
            WHERE id=?
        """, (
            result["answer"],
            json.dumps(result["citations"]),
            result["confidence"],
            result["evidence_snippet"],
            q["id"]
        ))
        results.append({"id": q["id"], **result})

    db.execute("UPDATE questionnaire_runs SET status='done' WHERE id=?", (run_id,))
    db.commit()
    return jsonify({"results": results})

@app.route("/run/<int:run_id>")
@login_required
def view_run(run_id):
    db = get_db()
    user = current_user()
    run = db.execute("SELECT * FROM questionnaire_runs WHERE id=?", (run_id,)).fetchone()
    if not run:
        flash("Run not found.", "error")
        return redirect(url_for("dashboard"))
    proj = db.execute("SELECT * FROM projects WHERE id=? AND user_id=?",
                      (run["project_id"], user["id"])).fetchone()
    if not proj:
        flash("Unauthorized.", "error")
        return redirect(url_for("dashboard"))
    questions = db.execute(
        "SELECT * FROM questions WHERE run_id=? ORDER BY question_number", (run_id,)
    ).fetchall()

    # Build summary stats
    total = len(questions)
    answered = sum(1 for q in questions if q["answer"] and q["answer"] != "Not found in references.")
    not_found = sum(1 for q in questions if q["answer"] == "Not found in references.")

    return render_template("run.html", user=user, project=proj, run=run,
                           questions=questions, total=total,
                           answered=answered, not_found=not_found)

@app.route("/run/<int:run_id>/edit_answer", methods=["POST"])
@login_required
def edit_answer(run_id):
    db = get_db()
    data = request.json
    q_id = data.get("question_id")
    edited = data.get("answer", "").strip()
    db.execute("UPDATE questions SET edited_answer=? WHERE id=? AND run_id=?",
               (edited, q_id, run_id))
    db.commit()
    return jsonify({"ok": True})

@app.route("/run/<int:run_id>/export")
@login_required
def export_run(run_id):
    db = get_db()
    user = current_user()
    run = db.execute("SELECT * FROM questionnaire_runs WHERE id=?", (run_id,)).fetchone()
    if not run:
        return "Not found", 404
    proj = db.execute("SELECT * FROM projects WHERE id=? AND user_id=?",
                      (run["project_id"], user["id"])).fetchone()
    if not proj:
        return "Unauthorized", 403

    questions = db.execute(
        "SELECT * FROM questions WHERE run_id=? ORDER BY question_number", (run_id,)
    ).fetchall()

    doc = Document()
    doc.core_properties.author = user["name"]

    # Title
    title = doc.add_heading(f"Questionnaire Answers — {proj['name']}", 0)
    title.runs[0].font.color.rgb = RGBColor(0x1a, 0x56, 0xDB)

    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Source file: {run['filename']}")
    doc.add_paragraph("")

    # Summary
    total = len(questions)
    answered = sum(1 for q in questions if q["answer"] and q["answer"] != "Not found in references.")
    not_found = total - answered
    summary = doc.add_paragraph()
    summary.add_run("Coverage Summary: ").bold = True
    summary.add_run(f"{total} questions | {answered} answered with citations | {not_found} not found in references")
    doc.add_paragraph("")

    for q in questions:
        # Question
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(f"Q{q['question_number']}. {q['question_text']}")
        q_run.bold = True
        q_run.font.size = Pt(11)

        # Answer (prefer edited)
        final_answer = q["edited_answer"] or q["answer"] or "Not found in references."
        a_para = doc.add_paragraph()
        a_para.add_run("Answer: ").bold = True
        a_para.add_run(final_answer)

        # Confidence
        if q["confidence"] and q["confidence"] > 0:
            conf_pct = int(float(q["confidence"]) * 100)
            c_para = doc.add_paragraph()
            c_run = c_para.add_run(f"Confidence: {conf_pct}%")
            c_run.italic = True
            c_run.font.size = Pt(9)

        # Citations
        if q["citations"]:
            try:
                cits = json.loads(q["citations"])
            except Exception:
                cits = [q["citations"]]
            if cits:
                cit_para = doc.add_paragraph()
                cit_run = cit_para.add_run("Citations: " + ", ".join(cits))
                cit_run.italic = True
                cit_run.font.size = Pt(9)
                cit_run.font.color.rgb = RGBColor(0x37, 0x74, 0xBB)

        # Evidence snippet
        if q["evidence_snippet"]:
            snip = q["evidence_snippet"][:300]
            s_para = doc.add_paragraph()
            s_run = s_para.add_run(f'Evidence: "{snip}..."')
            s_run.italic = True
            s_run.font.size = Pt(8)
            s_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        doc.add_paragraph("—" * 40)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    safe_name = re.sub(r"[^a-z0-9_-]", "_", proj["name"].lower())
    return send_file(
        buf,
        as_attachment=True,
        download_name=f"{safe_name}_answers.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@app.route("/run/<int:run_id>/delete", methods=["POST"])
@login_required
def delete_run(run_id):
    db = get_db()
    user = current_user()
    run = db.execute("SELECT * FROM questionnaire_runs WHERE id=?", (run_id,)).fetchone()
    if not run:
        return jsonify({"error": "Not found"}), 404
    proj = db.execute("SELECT * FROM projects WHERE id=? AND user_id=?",
                      (run["project_id"], user["id"])).fetchone()
    if not proj:
        return jsonify({"error": "Unauthorized"}), 403
    db.execute("DELETE FROM questions WHERE run_id=?", (run_id,))
    db.execute("DELETE FROM questionnaire_runs WHERE id=?", (run_id,))
    db.commit()
    return jsonify({"ok": True})

if __name__ == "__main__":
    app.run(debug=True, port=5000)